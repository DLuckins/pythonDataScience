# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx.shared import Cm
from matplotlib.backends.backend_pdf import PdfPages

df = pd.read_excel(r"dataset.xlsx")
fig, axes = plt.subplots(figsize=(16, 9), nrows=2, ncols=3)
df.plot.line(y="NewTime", ax=axes[0, 0], title="Line Plot Time")
df.plot(y="NewTime", title="Histogram Time", kind="hist", ax=axes[0, 1])
df.plot(y="NewTime", title="Cumulative Histogram Plot Time",
        kind="hist", cumulative=True, density=True, bins=100, histtype='step', ax=axes[0, 2])
df.plot.line(y="Total", title="Line Plot Total", ax=axes[1, 0])
df.plot(y="Total", title="Histogram Total", kind="hist", ax=axes[1, 1])
df.plot(y="Total", title="Cumulative Histogram Plot Total",
        kind="hist", cumulative=True, density=True, bins=100, histtype='step', ax=axes[1, 2])
plt.savefig(r"result1.png")
df.plot.bar(y="Rate", title="Bar plot of Rate variable", figsize=(16, 9))
fig.tight_layout()
plt.savefig(r"result2.png")
doc = docx.Document()
doc.add_heading("Practice 2, Deniss Lučkins st78712", 0)
doc.add_picture("result1.png", width=Cm(18))
doc.add_picture("result2.png", width=Cm(18))
stats1 = df["NewTime"].describe(percentiles=[0.9, 0.95, 0.99])
stats2 = df["Total"].describe(percentiles=[0.9, 0.95, 0.99])
stats3 = df["Rate"].describe(percentiles=[0.9, 0.95, 0.99])
n = len(df['NewTime'])
print(df)
output1 = ("Descriptive statistics for Time:\n" + str(stats1) + "\nSkewness:"
           + str(stats1.skew()) + "\nKurtosis:" + str(stats1.kurtosis()))
output2 = ("Descriptive statistics for Total:\n" + str(stats2) + "\nSkewness:"
           + str(stats2.skew()) + "\nKurtosis:" + str(stats2.kurtosis()))
output3 = ("Descriptive statistics for Rate:\n" + str(stats3) + "\nSkewness:"
           + str(stats3.skew()) + "\nKurtosis:" + str(stats3.kurtosis()))
doc.add_paragraph(output1)
doc.add_paragraph(output2)
doc.add_paragraph(output3)
print(output1)
print(output2)
print(output3)
print(df["Rate"].value_counts())
doc.save("result.docx")
